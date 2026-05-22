from datetime import datetime
from pathlib import Path
import argparse
from docx import Document


def add_heading(doc, text, level=1):
    try:
        doc.add_heading(text, level=level)
    except Exception:
        p = doc.add_paragraph()
        p.add_run(text)


def add_para(doc, text):
    doc.add_paragraph(text)


def add_figure_placeholder(doc, fig_no, title, content):
    add_para(doc, f"图{fig_no} {title}（待插入）")
    add_para(doc, f"图片内容建议：{content}")
    add_para(doc, "")


def main():
    parser = argparse.ArgumentParser(description="Generate thesis draft Word document from a template")
    parser.add_argument("--template", default="副本软件需求规格说明框架（公开）.docx", help="Template .docx file name in workspace root")
    parser.add_argument("--output-name", default="", help="Optional output file name (without path)")
    args = parser.parse_args()

    base_dir = Path(__file__).resolve().parents[1]
    template = base_dir / args.template
    if not template.exists():
        raise FileNotFoundError(f"Template not found: {template}")

    output_dir = base_dir / "output"
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d%H%M")
    if args.output_name:
        output_file = output_dir / args.output_name
    else:
        output_file = output_dir / f"论文草稿_Excel驱动Word自动生成系统_{timestamp}.docx"

    doc = Document(str(template))

    # Keep original template untouched and append thesis draft content at the end.
    doc.add_page_break()

    add_heading(doc, "Excel数据驱动Word文档自动生成系统设计与实现（论文草稿）", level=1)
    add_para(doc, "摘要：针对测试文档编制中人工复制粘贴效率低、易出错、格式难统一的问题，本文设计并实现了一个Excel数据驱动的Word文档自动生成系统。系统支持智能Sheet识别、模板样式继承、章节与表格动态编号、Web界面交互和REST API集成。实验结果表明，该系统可显著降低人工成本并提升文档规范性与可追溯性。")
    add_para(doc, "关键词：文档自动化；Excel；Word；模板填充；Spring Boot")

    add_heading(doc, "1 绪论", level=1)
    add_para(doc, "在软件测试文档场景中，测试大纲、测试报告和需求规格说明等文档通常包含大量结构化表格。传统人工填报流程耗时且易产生格式偏差。为此，本文提出并实现一种面向工程场景的自动化文档生成方案。")

    add_heading(doc, "2 需求分析", level=1)
    add_para(doc, "系统需支持四类Excel数据识别：测试用例、测试步骤、基本信息、列表型数据；需支持Word章节定位与Caption匹配填充；需提供命令行、Web和API三种使用方式。")

    add_heading(doc, "3 系统设计与实现", level=1)
    add_para(doc, "系统采用分层架构，包含读取层、处理层、文档生成层、接口层和存储层。通过配置化列名识别机制实现对异构Excel文件的兼容。文档生成阶段基于模板样式进行结构化插入，确保输出结果与模板风格一致。")

    add_heading(doc, "4 关键功能说明", level=1)
    add_para(doc, "（1）智能Sheet识别：不依赖固定Sheet名称，基于列特征判定数据类型。")
    add_para(doc, "（2）动态章节与表格编号：自动生成如5.2.1、5.2.2等子章节及对应表号。")
    add_para(doc, "（3）格式自适应：提取并复用模板中的字体、段落、表格格式。")
    add_para(doc, "（4）全流程服务化：支持上传、处理、预览、下载及历史管理。")

    add_heading(doc, "5 测试与结果分析", level=1)
    add_para(doc, "功能测试表明，系统可稳定完成多模块数据填充，目录可随内容动态更新。与人工方式相比，自动化方式在效率和一致性方面均表现更优。")

    add_heading(doc, "6 结论与展望", level=1)
    add_para(doc, "本文实现了一个可落地的Excel驱动Word自动生成系统。后续工作可在可视化规则配置、异常恢复策略、云存储与协同编辑能力方面继续扩展。")

    add_heading(doc, "7 待插入图片说明", level=1)
    add_figure_placeholder(doc, "1", "系统总体架构图", "展示输入层（Excel/模板）、处理层、生成层、存储层和CLI/Web/API入口关系。")
    add_figure_placeholder(doc, "2", "Excel智能识别流程图", "展示Sheet读取后按列特征分流识别为测试用例/测试步骤/基本信息/列表型数据。")
    add_figure_placeholder(doc, "3", "Word模板映射机制图", "展示Heading与Caption的定位规则，以及数据到表格单元格的映射路径。")
    add_figure_placeholder(doc, "4", "动态编号生成示意图", "展示5.2到5.2.1/5.2.2扩展及表5.2.1/表5.2.2自动生成过程。")
    add_figure_placeholder(doc, "5", "Web处理流程截图", "展示上传文件、开始处理、结果下载与历史管理页面关键区域。")
    add_figure_placeholder(doc, "6", "REST API时序图", "展示客户端调用/process、服务端处理、返回outputId并按ID下载的完整链路。")
    add_figure_placeholder(doc, "7", "性能对比图", "展示人工与自动化在耗时、错误率、一致性方面的柱状对比。")

    doc.save(str(output_file))
    print(str(output_file))


if __name__ == "__main__":
    main()
